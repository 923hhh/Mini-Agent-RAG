"""加载 Word、Excel、PPT 等 Office 文档内容。"""

from __future__ import annotations

import re
from collections.abc import Iterable
from pathlib import Path
from typing import Any

import docx2txt
from langchain_core.documents import Document

from .factory import BaseKnowledge


DOCX_HEADING_STYLE_PATTERN = re.compile(r"(?:heading|标题)\s*([1-3])", re.IGNORECASE)


class DocxKnowledge(BaseKnowledge):
    supported_extensions = (".docx",)

    def load(self) -> list[Document]:
        return _load_docx(self.path, self.base_metadata, self.relative_path)


class EpubKnowledge(BaseKnowledge):
    supported_extensions = (".epub",)

    def load(self) -> list[Document]:
        return _load_epub(self.path, self.base_metadata, self.relative_path)


def _load_docx(path: Path, base_metadata: dict[str, str], relative_path: str) -> list[Document]:
    try:
        from docx import Document as load_docx_document
    except ImportError as exc:
        raise RuntimeError(
            "解析 .docx 分节需要安装 `python-docx`。"
            "请在当前虚拟环境中执行 `pip install python-docx` 或 `pip install -r requirements.txt`。"
        ) from exc

    docx_document = load_docx_document(str(path))
    documents: list[Document] = []
    heading_stack: dict[int, str] = {}
    current_lines: list[str] = []
    current_section_title = str(base_metadata.get("title", path.stem))
    current_section_path = current_section_title
    current_section_index = 0

    def flush_section() -> None:
        nonlocal current_lines
        nonlocal current_section_index

        normalized_lines = [line.strip() for line in current_lines if line.strip()]
        if not normalized_lines:
            current_lines = []
            return

        section_title = current_section_title.strip() or path.stem
        section_path = current_section_path.strip() or section_title
        documents.append(
            Document(
                page_content="\n".join(normalized_lines),
                metadata={
                    **base_metadata,
                    "doc_id": f"{relative_path}#section-{current_section_index:04d}",
                    "title": section_title,
                    "section_title": section_title,
                    "section_path": section_path,
                    "section_index": current_section_index,
                },
            )
        )
        current_section_index += 1
        current_lines = []

    for paragraph in docx_document.paragraphs:
        text = _normalize_docx_text(paragraph.text)
        if not text:
            continue

        heading_level = _extract_docx_heading_level(paragraph)
        if heading_level is not None:
            flush_section()
            heading_stack = {
                level: title
                for level, title in heading_stack.items()
                if level < heading_level
            }
            heading_stack[heading_level] = text
            current_section_title = text
            current_section_path = " > ".join(
                title for level, title in sorted(heading_stack.items()) if title
            )
            current_lines = [text]
            continue

        current_lines.append(text)

    flush_section()
    if documents:
        return documents

    fallback_text = (docx2txt.process(str(path)) or "").strip()
    if not fallback_text:
        return []

    return [
        Document(
            page_content=fallback_text,
            metadata={
                **base_metadata,
                "doc_id": relative_path,
                "section_title": str(base_metadata.get("title", path.stem)),
                "section_path": str(base_metadata.get("title", path.stem)),
                "section_index": 0,
            },
        )
    ]


def _load_epub(path: Path, base_metadata: dict[str, str], relative_path: str) -> list[Document]:
    try:
        from ebooklib import ITEM_DOCUMENT, epub
    except ImportError as exc:
        raise RuntimeError(
            "解析 .epub 需要安装 `ebooklib`。"
            "请在当前虚拟环境中执行 `pip install ebooklib` 或 `pip install -r requirements.txt`。"
        ) from exc

    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError(
            "解析 .epub 需要安装 `beautifulsoup4`。"
            "请在当前虚拟环境中执行 `pip install beautifulsoup4` 或 `pip install -r requirements.txt`。"
        ) from exc

    book = epub.read_epub(str(path))
    book_title = _extract_epub_book_title(book) or str(base_metadata.get("title", path.stem))

    documents: list[Document] = []
    seen_item_ids: set[str] = set()
    spine_item_ids = _extract_epub_spine_item_ids(book.spine)

    for item_id in spine_item_ids:
        item = book.get_item_with_id(item_id)
        if item is None or item.get_type() != ITEM_DOCUMENT:
            continue
        section_document = _build_epub_document(
            item=item,
            base_metadata=base_metadata,
            relative_path=relative_path,
            book_title=book_title,
            section_index=len(documents),
            soup_builder=BeautifulSoup,
        )
        if section_document is None:
            continue
        documents.append(section_document)
        seen_item_ids.add(item_id)

    for item in book.get_items():
        if item.get_type() != ITEM_DOCUMENT:
            continue
        item_id = getattr(item, "id", "") or ""
        if item_id and item_id in seen_item_ids:
            continue
        section_document = _build_epub_document(
            item=item,
            base_metadata=base_metadata,
            relative_path=relative_path,
            book_title=book_title,
            section_index=len(documents),
            soup_builder=BeautifulSoup,
        )
        if section_document is None:
            continue
        documents.append(section_document)

    return documents


def _extract_docx_heading_level(paragraph) -> int | None:
    candidates: list[str] = []
    style = getattr(paragraph, "style", None)
    if style is not None:
        for attr_name in ("name", "style_id"):
            value = getattr(style, attr_name, "")
            if isinstance(value, str) and value.strip():
                candidates.append(value.strip())

    for candidate in candidates:
        match = DOCX_HEADING_STYLE_PATTERN.search(candidate)
        if not match:
            continue
        try:
            return int(match.group(1))
        except (TypeError, ValueError):
            return None
    return None


def _extract_epub_book_title(book: Any) -> str:
    try:
        metadata = book.get_metadata("DC", "title")
    except Exception:
        return ""
    if not metadata:
        return ""
    for item in metadata:
        if not item:
            continue
        raw_title = item[0] if isinstance(item, tuple) else item
        if isinstance(raw_title, str) and raw_title.strip():
            return raw_title.strip()
    return ""


def _extract_epub_spine_item_ids(spine: Iterable[Any]) -> list[str]:
    item_ids: list[str] = []
    for entry in spine:
        if isinstance(entry, tuple):
            item_id = str(entry[0]).strip()
        else:
            item_id = str(entry).strip()
        if not item_id or item_id == "nav":
            continue
        item_ids.append(item_id)
    return item_ids


def _build_epub_document(
    *,
    item: Any,
    base_metadata: dict[str, str],
    relative_path: str,
    book_title: str,
    section_index: int,
    soup_builder: Any,
) -> Document | None:
    try:
        raw_html = item.get_body_content()
    except Exception:
        raw_html = b""
    if not raw_html:
        return None

    soup = soup_builder(raw_html, "html.parser")
    _drop_epub_noise_nodes(soup)
    text = soup.get_text("\n", strip=True).strip()
    if not text:
        return None

    section_title = _extract_epub_section_title(soup) or f"{book_title} 第 {section_index + 1} 节"
    section_path = (
        book_title
        if section_title == book_title
        else f"{book_title} > {section_title}"
    )

    return Document(
        page_content=text,
        metadata={
            **base_metadata,
            "doc_id": f"{relative_path}#section-{section_index:04d}",
            "title": book_title,
            "section_title": section_title,
            "section_path": section_path,
            "section_index": section_index,
        },
    )


def _drop_epub_noise_nodes(soup: Any) -> None:
    for selector in ("nav", "script", "style"):
        for node in soup.find_all(selector):
            node.decompose()


def _extract_epub_section_title(soup: Any) -> str:
    for selector in ("h1", "h2", "h3", "title"):
        node = soup.find(selector)
        if node is None:
            continue
        text = node.get_text(" ", strip=True)
        if text:
            return text
    return ""


def _normalize_docx_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip()).strip()


__all__ = ["DocxKnowledge", "EpubKnowledge"]
