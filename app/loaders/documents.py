from __future__ import annotations

import os
import re
import shutil
import unicodedata
from collections.abc import Iterable
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Any

import docx2txt
from langchain_core.documents import Document
from pypdf import PdfReader

if TYPE_CHECKING:
    from app.services.settings import AppSettings


SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md"}
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".webp"}
DOCX_HEADING_STYLE_PATTERN = re.compile(r"(?:heading|标题)\s*([1-3])", re.IGNORECASE)


@dataclass
class PdfOutlineSection:
    title: str
    path: str
    level: int
    page_number: int
    has_children: bool = False


class BaseKnowledge(ABC):
    supported_extensions: tuple[str, ...] = ()

    def __init__(
        self,
        path: Path,
        content_dir: Path,
        settings: "AppSettings | None" = None,
    ) -> None:
        self.path = path
        self.content_dir = content_dir
        self.settings = settings
        self.relative_path = path.relative_to(content_dir).as_posix()
        self.base_metadata = {
            "source": path.name,
            "source_path": str(path.resolve()),
            "relative_path": self.relative_path,
            "extension": path.suffix.lower(),
            "title": path.stem,
            "content_type": "document_text",
            "source_modality": "text",
            "original_file_type": path.suffix.lower().lstrip("."),
            "evidence_summary": path.stem,
        }

    @classmethod
    def supports(cls, path: Path) -> bool:
        return path.suffix.lower() in cls.supported_extensions

    @abstractmethod
    def load(self) -> list[Document]:
        raise NotImplementedError


class TextKnowledge(BaseKnowledge):
    supported_extensions = (".txt",)

    def load(self) -> list[Document]:
        text = self.path.read_text(encoding="utf-8", errors="ignore")
        return [
            Document(
                page_content=text,
                metadata={
                    **self.base_metadata,
                    "doc_id": self.relative_path,
                },
            )
        ]


class MarkdownKnowledge(BaseKnowledge):
    supported_extensions = (".md",)

    def load(self) -> list[Document]:
        text = self.path.read_text(encoding="utf-8", errors="ignore")
        return [
            Document(
                page_content=text,
                metadata={
                    **self.base_metadata,
                    "doc_id": self.relative_path,
                },
            )
        ]


class PdfKnowledge(BaseKnowledge):
    supported_extensions = (".pdf",)

    def load(self) -> list[Document]:
        return _load_pdf(self.path, self.base_metadata, self.relative_path)


class DocxKnowledge(BaseKnowledge):
    supported_extensions = (".docx",)

    def load(self) -> list[Document]:
        return _load_docx(self.path, self.base_metadata, self.relative_path)


class EpubKnowledge(BaseKnowledge):
    supported_extensions = (".epub",)

    def load(self) -> list[Document]:
        return _load_epub(self.path, self.base_metadata, self.relative_path)


class ImageKnowledge(BaseKnowledge):
    supported_extensions = tuple(sorted(SUPPORTED_IMAGE_EXTENSIONS))

    def load(self) -> list[Document]:
        return _load_image(
            self.path,
            self.base_metadata,
            self.relative_path,
            settings=self.settings,
        )


class KnowledgeFactory:
    _REGISTRY = [
        MarkdownKnowledge,
        TextKnowledge,
        PdfKnowledge,
        DocxKnowledge,
        EpubKnowledge,
        ImageKnowledge,
    ]

    @classmethod
    def create(
        cls,
        path: Path,
        content_dir: Path,
        settings: "AppSettings | None" = None,
    ) -> BaseKnowledge:
        for knowledge_cls in cls._REGISTRY:
            if knowledge_cls.supports(path):
                return knowledge_cls(path, content_dir, settings=settings)
        raise ValueError(f"暂不支持的文件类型: {path}")

    @classmethod
    def load(
        cls,
        path: Path,
        content_dir: Path,
        settings: "AppSettings | None" = None,
    ) -> list[Document]:
        return cls.create(path, content_dir, settings=settings).load()


def list_supported_files(content_dir: Path, supported_extensions: list[str]) -> list[Path]:
    normalized = {ext.lower() for ext in supported_extensions}
    files: list[Path] = []
    for path in content_dir.rglob("*"):
        if path.is_file() and path.suffix.lower() in normalized:
            files.append(path)
    return sorted(files)


def load_documents(
    content_dir: Path,
    supported_extensions: list[str],
    settings: "AppSettings | None" = None,
) -> tuple[list[Document], list[Path]]:
    files = list_supported_files(content_dir, supported_extensions)
    documents: list[Document] = []

    for path in files:
        documents.extend(load_file(path, content_dir, settings=settings))

    return documents, files


def load_file(
    path: Path,
    content_dir: Path,
    settings: "AppSettings | None" = None,
) -> list[Document]:
    return KnowledgeFactory.load(path, content_dir, settings=settings)


def _load_pdf(path: Path, base_metadata: dict[str, str], relative_path: str) -> list[Document]:
    reader = PdfReader(str(path))
    outlined_documents = _load_pdf_outline_sections(reader, base_metadata, relative_path)
    if outlined_documents:
        return outlined_documents

    documents: list[Document] = []

    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={
                    **base_metadata,
                    "doc_id": f"{relative_path}#page-{index}",
                    "page": index,
                    "page_end": index,
                    "section_title": f"{base_metadata.get('title', path.stem)} 第 {index} 页",
                    "section_path": f"{base_metadata.get('title', path.stem)} > 第 {index} 页",
                    "section_index": index - 1,
                },
            )
        )

    return documents


def _load_pdf_outline_sections(
    reader: PdfReader,
    base_metadata: dict[str, str],
    relative_path: str,
) -> list[Document]:
    outline_root = getattr(reader, "outline", None)
    if not outline_root:
        return []

    sections = _flatten_pdf_outline(reader, outline_root)
    if not sections:
        return []

    total_pages = len(reader.pages)
    documents: list[Document] = []
    for index, section in enumerate(sections):
        next_section_page = _find_next_distinct_section_page(sections, index)
        if section.has_children and next_section_page == section.page_number:
            continue

        start_page = max(1, min(section.page_number, total_pages))
        end_page = total_pages
        if next_section_page is not None and next_section_page > start_page:
            end_page = min(total_pages, next_section_page - 1)
        elif next_section_page is not None and next_section_page <= start_page:
            end_page = start_page

        page_texts: list[str] = []
        for page_number in range(start_page, end_page + 1):
            text = reader.pages[page_number - 1].extract_text() or ""
            normalized = text.strip()
            if normalized:
                page_texts.append(normalized)
        if not page_texts:
            continue

        documents.append(
            Document(
                page_content="\n\n".join(page_texts),
                metadata={
                    **base_metadata,
                    "doc_id": f"{relative_path}#section-{len(documents):04d}",
                    "page": start_page,
                    "page_end": end_page,
                    "title": section.title,
                    "section_title": section.title,
                    "section_path": section.path,
                    "section_index": len(documents),
                },
            )
        )

    return documents


def _load_docx(path: Path, base_metadata: dict[str, str], relative_path: str) -> list[Document]:
    try:
        from docx import Document as load_docx_document
    except ImportError as exc:
        raise RuntimeError(
            "解析 .docx 分节需要安装 `python-docx`。"
            "请在当前虚拟环境中执行 `pip install python-docx` 或 `pip install -r requirements.txt`。"
        ) from exc

    docx_document = load_docx_document(str(path))
    documents: list[Document] = []
    heading_stack: dict[int, str] = {}
    current_lines: list[str] = []
    current_section_title = str(base_metadata.get("title", path.stem))
    current_section_path = current_section_title
    current_section_index = 0

    def flush_section() -> None:
        nonlocal current_lines
        nonlocal current_section_index

        normalized_lines = [line.strip() for line in current_lines if line.strip()]
        if not normalized_lines:
            current_lines = []
            return

        section_title = current_section_title.strip() or path.stem
        section_path = current_section_path.strip() or section_title
        documents.append(
            Document(
                page_content="\n".join(normalized_lines),
                metadata={
                    **base_metadata,
                    "doc_id": f"{relative_path}#section-{current_section_index:04d}",
                    "title": section_title,
                    "section_title": section_title,
                    "section_path": section_path,
                    "section_index": current_section_index,
                },
            )
        )
        current_section_index += 1
        current_lines = []

    for paragraph in docx_document.paragraphs:
        text = _normalize_docx_text(paragraph.text)
        if not text:
            continue

        heading_level = _extract_docx_heading_level(paragraph)
        if heading_level is not None:
            flush_section()
            heading_stack = {
                level: title
                for level, title in heading_stack.items()
                if level < heading_level
            }
            heading_stack[heading_level] = text
            current_section_title = text
            current_section_path = " > ".join(
                title for level, title in sorted(heading_stack.items()) if title
            )
            current_lines = [text]
            continue

        current_lines.append(text)

    flush_section()
    if documents:
        return documents

    fallback_text = (docx2txt.process(str(path)) or "").strip()
    if not fallback_text:
        return []

    return [
        Document(
            page_content=fallback_text,
            metadata={
                **base_metadata,
                "doc_id": relative_path,
                "section_title": str(base_metadata.get("title", path.stem)),
                "section_path": str(base_metadata.get("title", path.stem)),
                "section_index": 0,
            },
        )
    ]


def _load_epub(path: Path, base_metadata: dict[str, str], relative_path: str) -> list[Document]:
    try:
        from ebooklib import ITEM_DOCUMENT, epub
    except ImportError as exc:
        raise RuntimeError(
            "解析 .epub 需要安装 `ebooklib`。"
            "请在当前虚拟环境中执行 `pip install ebooklib` 或 `pip install -r requirements.txt`。"
        ) from exc

    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError(
            "解析 .epub 需要安装 `beautifulsoup4`。"
            "请在当前虚拟环境中执行 `pip install beautifulsoup4` 或 `pip install -r requirements.txt`。"
        ) from exc

    book = epub.read_epub(str(path))
    book_title = _extract_epub_book_title(book) or str(base_metadata.get("title", path.stem))

    documents: list[Document] = []
    seen_item_ids: set[str] = set()
    spine_item_ids = _extract_epub_spine_item_ids(book.spine)

    for item_id in spine_item_ids:
        item = book.get_item_with_id(item_id)
        if item is None or item.get_type() != ITEM_DOCUMENT:
            continue
        section_document = _build_epub_document(
            item=item,
            base_metadata=base_metadata,
            relative_path=relative_path,
            book_title=book_title,
            section_index=len(documents),
            soup_builder=BeautifulSoup,
        )
        if section_document is None:
            continue
        documents.append(section_document)
        seen_item_ids.add(item_id)

    for item in book.get_items():
        if item.get_type() != ITEM_DOCUMENT:
            continue
        item_id = getattr(item, "id", "") or ""
        if item_id and item_id in seen_item_ids:
            continue
        section_document = _build_epub_document(
            item=item,
            base_metadata=base_metadata,
            relative_path=relative_path,
            book_title=book_title,
            section_index=len(documents),
            soup_builder=BeautifulSoup,
        )
        if section_document is None:
            continue
        documents.append(section_document)

    return documents


def _load_image(
    path: Path,
    base_metadata: dict[str, str],
    relative_path: str,
    *,
    settings: "AppSettings | None" = None,
) -> list[Document]:
    ocr_enabled = True if settings is None else settings.kb.IMAGE_OCR_ENABLED
    image_title = base_metadata.get("title", path.stem)
    ocr_text = ""
    caption_text = ""
    ocr_error = ""
    caption_error = ""
    image_caption_reason = "disabled"
    ocr_char_count = 0

    if ocr_enabled:
        language = "chi_sim+eng" if settings is None else settings.kb.IMAGE_OCR_LANGUAGE
        if settings is None:
            tesseract_cmd = ""
        else:
            configured_cmd = settings.kb.OCR_TESSERACT_CMD.strip()
            tesseract_cmd = (
                str(settings.resolve_path(configured_cmd))
                if configured_cmd
                else ""
            )
        try:
            ocr_text = _extract_image_text(
                path,
                language=language,
                tesseract_cmd=tesseract_cmd,
                min_confidence=(
                    60.0 if settings is None else settings.kb.OCR_MIN_CONFIDENCE
                ),
                min_text_length=(
                    6 if settings is None else settings.kb.OCR_MIN_TEXT_LENGTH
                ),
                min_meaningful_ratio=(
                    0.6 if settings is None else settings.kb.OCR_MIN_MEANINGFUL_RATIO
                ),
            ).strip()
        except Exception as exc:
            ocr_error = str(exc)

    should_caption, image_caption_reason, ocr_char_count = _should_generate_image_caption(
        settings=settings,
        ocr_text=ocr_text,
    )
    if should_caption:
        try:
            from app.services.image_caption_service import caption_image
            raw_caption_text = caption_image(settings, path, ocr_text=ocr_text).strip()
            if _is_low_quality_image_caption(raw_caption_text):
                caption_text = ""
                caption_error = "low_quality_caption_filtered"
            else:
                caption_text = raw_caption_text
        except Exception as exc:
            caption_error = str(exc)

    normalized = _build_image_knowledge_text(
        path=path,
        ocr_text=ocr_text,
        caption_text=caption_text,
        ocr_enabled=ocr_enabled,
        ocr_error=ocr_error,
        caption_enabled=should_caption,
        caption_error=caption_error,
    )

    return [
        Document(
            page_content=normalized,
            metadata={
                **base_metadata,
                "doc_id": relative_path,
                "section_title": image_title,
                "section_path": image_title,
                "section_index": 0,
                "content_type": "image_evidence",
                "source_modality": _resolve_image_source_modality(
                    ocr_text=ocr_text,
                    caption_text=caption_text,
                ),
                "ocr_text": ocr_text or None,
                "image_caption": caption_text or None,
                "evidence_summary": _build_image_evidence_summary(
                    image_title=image_title,
                    ocr_text=ocr_text,
                    caption_text=caption_text,
                ),
                "image_ocr_enabled": ocr_enabled,
                "image_caption_enabled": should_caption,
                "image_caption_reason": image_caption_reason,
                "ocr_char_count": ocr_char_count,
                "ocr_filtered_char_count": len(re.sub(r"\s+", "", ocr_text)),
                "ocr_language": (
                    "chi_sim+eng" if settings is None else settings.kb.IMAGE_OCR_LANGUAGE
                ) if ocr_enabled else "",
                "image_caption_model": (
                    settings.model.IMAGE_VLM_MODEL if settings is not None and should_caption else ""
                ),
                "ocr_error": ocr_error,
                "image_caption_error": caption_error,
            },
        )
    ]


def _flatten_pdf_outline(
    reader: PdfReader,
    items: list[Any],
    *,
    level: int = 1,
    parent_path: list[str] | None = None,
) -> list[PdfOutlineSection]:
    base_path = list(parent_path or [])
    flattened: list[PdfOutlineSection] = []
    last_section: PdfOutlineSection | None = None

    for item in items:
        if isinstance(item, list):
            child_parent = base_path
            if last_section is not None:
                last_section.has_children = True
                child_parent = last_section.path.split(" > ")
            flattened.extend(
                _flatten_pdf_outline(
                    reader,
                    item,
                    level=level + 1,
                    parent_path=child_parent,
                )
            )
            continue

        title = _extract_pdf_outline_title(item)
        if not title:
            continue
        try:
            page_number = int(reader.get_destination_page_number(item)) + 1
        except Exception:
            continue

        path_parts = [part for part in (*base_path, title) if part]
        section = PdfOutlineSection(
            title=title,
            path=" > ".join(path_parts),
            level=level,
            page_number=page_number,
        )
        flattened.append(section)
        last_section = section

    return flattened


def _extract_pdf_outline_title(item: Any) -> str:
    for attr_name in ("title", "/Title"):
        value = getattr(item, attr_name, None)
        if isinstance(value, str) and value.strip():
            return value.strip()
        if isinstance(item, dict):
            dict_value = item.get(attr_name)
            if isinstance(dict_value, str) and dict_value.strip():
                return dict_value.strip()
    return ""


def _find_next_distinct_section_page(
    sections: list[PdfOutlineSection],
    current_index: int,
) -> int | None:
    current_page = sections[current_index].page_number
    for section in sections[current_index + 1 :]:
        if section.page_number != current_page:
            return section.page_number
    return None


def _extract_docx_heading_level(paragraph) -> int | None:
    candidates: list[str] = []
    style = getattr(paragraph, "style", None)
    if style is not None:
        for attr_name in ("name", "style_id"):
            value = getattr(style, attr_name, "")
            if isinstance(value, str) and value.strip():
                candidates.append(value.strip())

    for candidate in candidates:
        match = DOCX_HEADING_STYLE_PATTERN.search(candidate)
        if not match:
            continue
        try:
            return int(match.group(1))
        except (TypeError, ValueError):
            return None
    return None


def _extract_epub_book_title(book: Any) -> str:
    try:
        metadata = book.get_metadata("DC", "title")
    except Exception:
        return ""
    if not metadata:
        return ""
    for item in metadata:
        if not item:
            continue
        raw_title = item[0] if isinstance(item, tuple) else item
        if isinstance(raw_title, str) and raw_title.strip():
            return raw_title.strip()
    return ""


def _extract_epub_spine_item_ids(spine: Iterable[Any]) -> list[str]:
    item_ids: list[str] = []
    for entry in spine:
        if isinstance(entry, tuple):
            item_id = str(entry[0]).strip()
        else:
            item_id = str(entry).strip()
        if not item_id or item_id == "nav":
            continue
        item_ids.append(item_id)
    return item_ids


def _build_epub_document(
    *,
    item: Any,
    base_metadata: dict[str, str],
    relative_path: str,
    book_title: str,
    section_index: int,
    soup_builder: Any,
) -> Document | None:
    try:
        raw_html = item.get_body_content()
    except Exception:
        raw_html = b""
    if not raw_html:
        return None

    soup = soup_builder(raw_html, "html.parser")
    _drop_epub_noise_nodes(soup)
    text = soup.get_text("\n", strip=True).strip()
    if not text:
        return None

    section_title = _extract_epub_section_title(soup) or f"{book_title} 第 {section_index + 1} 节"
    section_path = (
        book_title
        if section_title == book_title
        else f"{book_title} > {section_title}"
    )

    return Document(
        page_content=text,
        metadata={
            **base_metadata,
            "doc_id": f"{relative_path}#section-{section_index:04d}",
            "title": book_title,
            "section_title": section_title,
            "section_path": section_path,
            "section_index": section_index,
        },
    )


def _drop_epub_noise_nodes(soup: Any) -> None:
    for selector in ("nav", "script", "style"):
        for node in soup.find_all(selector):
            node.decompose()


def _extract_epub_section_title(soup: Any) -> str:
    for selector in ("h1", "h2", "h3", "title"):
        node = soup.find(selector)
        if node is None:
            continue
        text = node.get_text(" ", strip=True)
        if text:
            return text
    return ""


def _normalize_docx_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip()).strip()


def _extract_image_text(
    path: Path,
    *,
    language: str,
    tesseract_cmd: str,
    min_confidence: float,
    min_text_length: int,
    min_meaningful_ratio: float,
) -> str:
    try:
        import pytesseract
    except ImportError as exc:
        raise RuntimeError(
            "图片 OCR 需要安装 `pytesseract`。"
            "请在当前虚拟环境中执行 `pip install pytesseract` 或 `pip install -r requirements.txt`。"
        ) from exc

    try:
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError(
            "图片 OCR 需要安装 `Pillow`。"
            "请在当前虚拟环境中执行 `pip install Pillow` 或 `pip install -r requirements.txt`。"
        ) from exc

    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        tessdata_dir = Path(tesseract_cmd).resolve().parent / "tessdata"
        if tessdata_dir.exists():
            os.environ["TESSDATA_PREFIX"] = str(tessdata_dir)
    elif not shutil.which("tesseract"):
        raise RuntimeError(
            "图片 OCR 未检测到 `tesseract` 可执行文件。"
            "请先安装 Tesseract OCR，或在 kb_settings.yaml 中配置 OCR_TESSERACT_CMD。"
        )

    collected_lines: list[str] = []
    seen_lines: set[str] = set()
    with Image.open(path) as original_image:
        for candidate_image in _generate_ocr_candidate_images(original_image):
            data = pytesseract.image_to_data(
                candidate_image,
                lang=language,
                output_type=pytesseract.Output.DICT,
            )
            for line_text in _extract_accepted_ocr_lines(
                data,
                min_confidence=min_confidence,
                min_text_length=min_text_length,
                min_meaningful_ratio=min_meaningful_ratio,
            ):
                normalized = line_text.strip()
                if not normalized or normalized in seen_lines:
                    continue
                seen_lines.add(normalized)
                collected_lines.append(normalized)

    return "\n".join(collected_lines).strip()


def _extract_accepted_ocr_lines(
    data: dict[str, Any],
    *,
    min_confidence: float,
    min_text_length: int,
    min_meaningful_ratio: float,
) -> list[str]:
    text_items = data.get("text") or []
    conf_items = data.get("conf") or []
    block_items = data.get("block_num") or []
    par_items = data.get("par_num") or []
    line_items = data.get("line_num") or []

    line_tokens: dict[tuple[int, int, int], list[str]] = {}
    for index, raw_text in enumerate(text_items):
        token = _normalize_ocr_token(raw_text)
        if not token:
            continue

        raw_confidence = conf_items[index] if index < len(conf_items) else None
        confidence = _parse_ocr_confidence(raw_confidence)
        token_min_confidence = _effective_ocr_confidence_threshold(
            token,
            base_min_confidence=min_confidence,
        )
        if confidence is None or confidence < token_min_confidence:
            continue
        if not _is_meaningful_ocr_token(token):
            continue

        key = (
            int(block_items[index] if index < len(block_items) else 1 or 1),
            int(par_items[index] if index < len(par_items) else 1 or 1),
            int(line_items[index] if index < len(line_items) else 1 or 1),
        )
        line_tokens.setdefault(key, []).append(token)

    accepted_lines: list[str] = []
    for tokens in line_tokens.values():
        line_text = _join_ocr_tokens(tokens)
        if _should_keep_ocr_line(
            line_text,
            tokens=tokens,
            min_text_length=min_text_length,
            min_meaningful_ratio=min_meaningful_ratio,
        ):
            accepted_lines.append(line_text)
    return accepted_lines


def _generate_ocr_candidate_images(image) -> list[Any]:
    try:
        from PIL import ImageFilter, ImageOps
    except ImportError:
        return [image.copy()]

    base = image.convert("RGB")
    variants: list[Any] = [base]

    grayscale = ImageOps.grayscale(base)
    variants.append(ImageOps.autocontrast(grayscale))

    enlarged = grayscale.resize(
        (max(1, grayscale.width * 2), max(1, grayscale.height * 2))
    )
    variants.append(ImageOps.autocontrast(enlarged))

    sharpened = ImageOps.autocontrast(grayscale.filter(ImageFilter.SHARPEN))
    variants.append(sharpened)

    thresholded = ImageOps.autocontrast(grayscale).point(
        lambda value: 255 if value > 160 else 0
    )
    variants.append(thresholded)

    crop_boxes = _build_ocr_crop_boxes(base.width, base.height)
    if crop_boxes:
        for left, top, right, bottom in crop_boxes:
            cropped = base.crop((left, top, right, bottom))
            variants.append(cropped)
            cropped_gray = ImageOps.autocontrast(ImageOps.grayscale(cropped))
            variants.append(cropped_gray)

    return variants


def _build_ocr_crop_boxes(width: int, height: int) -> list[tuple[int, int, int, int]]:
    if width < 240 or height < 120:
        return []

    boxes: list[tuple[int, int, int, int]] = []
    top_cut = max(1, int(height * 0.4))
    bottom_start = max(0, int(height * 0.6))
    middle_top = max(0, int(height * 0.2))
    middle_bottom = min(height, int(height * 0.8))

    boxes.append((0, 0, width, top_cut))
    boxes.append((0, bottom_start, width, height))
    boxes.append((0, middle_top, width, middle_bottom))

    if width >= 320:
        left_cut = max(1, int(width * 0.55))
        right_start = max(0, int(width * 0.45))
        boxes.append((0, 0, left_cut, height))
        boxes.append((right_start, 0, width, height))

    deduped: list[tuple[int, int, int, int]] = []
    seen: set[tuple[int, int, int, int]] = set()
    for box in boxes:
        if box in seen:
            continue
        seen.add(box)
        deduped.append(box)
    return deduped


def _should_generate_image_caption(
    *,
    settings: "AppSettings | None",
    ocr_text: str,
) -> tuple[bool, str, int]:
    if settings is None:
        return False, "settings_missing", 0
    if not settings.model.IMAGE_VLM_ENABLED:
        return False, "vlm_disabled", 0

    normalized_ocr_text = _normalize_ocr_text_for_caption(ocr_text)
    ocr_char_count = len(normalized_ocr_text)

    if settings.model.IMAGE_VLM_AUTO_TRIGGER_BY_OCR:
        threshold = settings.model.IMAGE_VLM_SKIP_IF_OCR_CHARS_AT_LEAST
        if ocr_char_count >= threshold:
            return False, "ocr_rich_skip_vlm", ocr_char_count
        return True, "ocr_sparse_use_vlm", ocr_char_count

    if settings.model.IMAGE_VLM_ONLY_WHEN_OCR_EMPTY:
        if ocr_char_count == 0:
            return True, "ocr_empty_use_vlm", ocr_char_count
        return False, "ocr_non_empty_skip_vlm", ocr_char_count

    return True, "always_use_vlm", ocr_char_count


def _normalize_ocr_text_for_caption(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _normalize_ocr_token(text: str) -> str:
    return re.sub(r"\s+", "", text or "").strip()


def _parse_ocr_confidence(raw_value: Any) -> float | None:
    try:
        value = float(str(raw_value).strip())
    except (TypeError, ValueError):
        return None
    if value < 0:
        return None
    return value


def _join_ocr_tokens(tokens: list[str]) -> str:
    if not tokens:
        return ""

    parts: list[str] = [tokens[0]]
    for token in tokens[1:]:
        previous_last = parts[-1][-1]
        current_first = token[0]
        if _should_insert_space_between(previous_last, current_first):
            parts.append(f" {token}")
        else:
            parts.append(token)
    return "".join(parts).strip()


def _should_insert_space_between(left_char: str, right_char: str) -> bool:
    return _is_ascii_word_char(left_char) and _is_ascii_word_char(right_char)


def _is_ascii_word_char(char: str) -> bool:
    return char.isascii() and char.isalnum()


def _is_meaningful_ocr_token(token: str) -> bool:
    compact = re.sub(r"\s+", "", token)
    if not compact:
        return False
    meaningful_count = sum(1 for char in compact if _is_meaningful_ocr_char(char))
    return meaningful_count >= max(1, len(compact) // 2)


def _should_keep_ocr_line(
    text: str,
    *,
    tokens: list[str],
    min_text_length: int,
    min_meaningful_ratio: float,
) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if not compact:
        return False

    contains_cjk = any(_is_cjk_char(char) for char in compact)
    contains_ascii_letters = any(char.isascii() and char.isalpha() for char in compact)

    effective_min_length = _effective_min_ocr_line_length(
        contains_cjk=contains_cjk,
        contains_ascii_letters=contains_ascii_letters,
        base_min_text_length=min_text_length,
    )
    if len(compact) < effective_min_length:
        return False

    meaningful_count = sum(1 for char in compact if _is_meaningful_ocr_char(char))
    effective_min_ratio = _effective_min_meaningful_ratio(
        contains_cjk=contains_cjk,
        contains_ascii_letters=contains_ascii_letters,
        base_min_meaningful_ratio=min_meaningful_ratio,
    )
    if meaningful_count / len(compact) < effective_min_ratio:
        return False
    if _looks_like_short_mixed_script_noise(tokens):
        return False
    if _looks_like_mixed_script_noise_line(compact):
        return False
    if _looks_like_latin_short_token_noise(tokens):
        return False
    if _looks_like_repeated_symbol_noise(compact):
        return False
    return True


def _looks_like_short_mixed_script_noise(tokens: list[str]) -> bool:
    if len(tokens) < 3:
        return False
    if not all(len(token) <= 3 for token in tokens):
        return False

    has_cjk = any(any(_is_cjk_char(char) for char in token) for token in tokens)
    has_latin = any(any(char.isascii() and char.isalpha() for char in token) for token in tokens)
    return has_cjk and has_latin


def _looks_like_mixed_script_noise_line(text: str) -> bool:
    has_cjk = any(_is_cjk_char(char) for char in text)
    has_latin = any(char.isascii() and char.isalpha() for char in text)
    if not (has_cjk and has_latin):
        return False
    return len(text) < 12


def _looks_like_latin_short_token_noise(tokens: list[str]) -> bool:
    latin_tokens = [
        token for token in tokens if any(char.isascii() and char.isalpha() for char in token)
    ]
    if not latin_tokens:
        return False
    if len(latin_tokens) >= 3 and all(len(token) <= 3 for token in latin_tokens):
        return True
    average_length = sum(len(token) for token in latin_tokens) / len(latin_tokens)
    return average_length < 2.2


def _looks_like_repeated_symbol_noise(text: str) -> bool:
    if not text:
        return True
    symbol_count = sum(1 for char in text if not _is_meaningful_ocr_char(char))
    return symbol_count > 0 and symbol_count / len(text) > 0.35


def _effective_ocr_confidence_threshold(
    token: str,
    *,
    base_min_confidence: float,
) -> float:
    compact = re.sub(r"\s+", "", token)
    if not compact:
        return base_min_confidence

    has_cjk = any(_is_cjk_char(char) for char in compact)
    has_latin = any(char.isascii() and char.isalpha() for char in compact)
    if has_cjk:
        if len(compact) >= 2:
            return max(40.0, base_min_confidence - 20.0)
        return max(50.0, base_min_confidence - 10.0)
    if has_latin and len(compact) <= 2:
        return min(95.0, base_min_confidence + 15.0)
    if has_latin and len(compact) <= 4:
        return min(90.0, base_min_confidence + 5.0)
    return base_min_confidence


def _effective_min_ocr_line_length(
    *,
    contains_cjk: bool,
    contains_ascii_letters: bool,
    base_min_text_length: int,
) -> int:
    if contains_cjk:
        return max(4, base_min_text_length - 2)
    if contains_ascii_letters:
        return max(8, base_min_text_length + 2)
    return base_min_text_length


def _effective_min_meaningful_ratio(
    *,
    contains_cjk: bool,
    contains_ascii_letters: bool,
    base_min_meaningful_ratio: float,
) -> float:
    if contains_cjk:
        return min(base_min_meaningful_ratio, 0.5)
    if contains_ascii_letters:
        return max(base_min_meaningful_ratio, 0.75)
    return base_min_meaningful_ratio


def _is_meaningful_ocr_char(char: str) -> bool:
    if _is_cjk_char(char):
        return True
    if char.isdigit():
        return True
    if char.isascii() and char.isalpha():
        return True

    category = unicodedata.category(char)
    return category.startswith("L") or category == "Nd"


def _is_cjk_char(char: str) -> bool:
    codepoint = ord(char)
    return (
        0x4E00 <= codepoint <= 0x9FFF
        or 0x3400 <= codepoint <= 0x4DBF
        or 0xF900 <= codepoint <= 0xFAFF
    )


def _build_image_knowledge_text(
    *,
    path: Path,
    ocr_text: str,
    caption_text: str,
    ocr_enabled: bool,
    ocr_error: str,
    caption_enabled: bool,
    caption_error: str,
) -> str:
    sections: list[str] = [f"图片文件: {path.name}"]

    if ocr_text:
        sections.append(f"[图片文字 OCR]\n{ocr_text}")
    elif ocr_enabled and ocr_error:
        sections.append(f"[图片文字 OCR 失败]\n{ocr_error}")
    elif ocr_enabled:
        sections.append("[图片文字 OCR]\n未识别到明显文字。")
    else:
        sections.append("[图片文字 OCR]\n已关闭。")

    if caption_text:
        sections.append(f"[图片内容描述]\n{caption_text}")
    elif caption_enabled and caption_error:
        sections.append(f"[图片内容描述生成失败]\n{caption_error}")
    elif caption_enabled:
        sections.append("[图片内容描述]\n未生成有效描述。")

    return "\n\n".join(section.strip() for section in sections if section.strip())


def _resolve_image_source_modality(
    *,
    ocr_text: str,
    caption_text: str,
) -> str:
    if ocr_text and caption_text:
        return "ocr+vision"
    if ocr_text:
        return "ocr"
    if caption_text:
        return "vision"
    return "image"


def _build_image_evidence_summary(
    *,
    image_title: str,
    ocr_text: str,
    caption_text: str,
) -> str:
    parts = [image_title]
    if ocr_text:
        compact_ocr = re.sub(r"\s+", " ", ocr_text).strip()
        parts.append(f"OCR: {compact_ocr[:80]}")
    if caption_text:
        compact_caption = re.sub(r"\s+", " ", caption_text).strip()
        parts.append(f"Vision: {compact_caption[:80]}")
    return " | ".join(part for part in parts if part).strip()


def _is_low_quality_image_caption(text: str) -> bool:
    normalized = re.sub(r"\s+", "", text or "").lower()
    if not normalized:
        return True

    blocked_patterns = (
        "无法为你提供相应解答",
        "你可以尝试提供其他话题",
        "我会尽力为你提供支持和解答",
        "抱歉",
        "无法判断",
        "无法直接判断",
        "依据不足",
        "无法确认",
        "未提供足够信息",
        "不能为你提供",
    )
    if any(pattern in normalized for pattern in blocked_patterns):
        return True

    if len(normalized) < 8:
        return True
    return False
